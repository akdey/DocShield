from pathlib import Path
import docx
from .base import BaseParser, ParsedDocument
import re

class WordParser(BaseParser):
    def read(self, path: Path) -> ParsedDocument:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # We can also extract tables if needed later.
        text = "\n".join(full_text)
        return ParsedDocument(original_path=path, text=text)

    def write_masked(self, original_path: Path, output_path: Path, text: str, replacements: list[tuple[int, int, str]]) -> None:
        """
        Replacing text in DOCX while preserving formatting is complex because of 'runs'.
        We use the exact character coordinates (replacements) provided by the scanner
        to modify the specific runs in place.
        """
        doc = docx.Document(original_path)
        
        # We need to map global character indices to paragraph and run.
        # Let's compute paragraph bounds
        para_bounds = []
        current_offset = 0
        for para in doc.paragraphs:
            para_len = len(para.text)
            para_bounds.append((para, current_offset, current_offset + para_len))
            current_offset += para_len + 1 # +1 for the \n added during join
            
        # Group replacements by paragraph
        # Because we replace text inline, we must process them backwards 
        # so earlier indices don't shift!
        replacements.sort(key=lambda x: x[0], reverse=True)
        
        for rep_start, rep_end, _, rep_text in replacements:
            # Find which paragraph this replacement belongs to
            for para, p_start, p_end in para_bounds:
                if p_start <= rep_start < p_end:
                    local_start = rep_start - p_start
                    local_end = rep_end - p_start
                    self._replace_in_runs(para, local_start, local_end, rep_text)
                    break
                    
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _replace_in_runs(self, para, local_start: int, local_end: int, replacement: str):
        # Map each character in para.text to a (run_index, char_index_in_run)
        run_map = []
        for r_idx, run in enumerate(para.runs):
            for c_idx, char in enumerate(run.text):
                run_map.append((r_idx, c_idx))
                
        if local_start >= len(run_map):
            return
            
        if local_end > len(run_map):
            local_end = len(run_map)

        first_run_idx, first_char_idx = run_map[local_start]
        last_run_idx, last_char_idx = run_map[local_end - 1]

        # Rebuild the text for each affected run
        for r_idx in range(first_run_idx, last_run_idx + 1):
            run = para.runs[r_idx]
            run_text = list(run.text)
            
            s_idx = first_char_idx if r_idx == first_run_idx else 0
            e_idx = last_char_idx + 1 if r_idx == last_run_idx else len(run.text)
            
            if r_idx == first_run_idx:
                run_text[s_idx:e_idx] = list(replacement)
            else:
                run_text[s_idx:e_idx] = []
                
            run.text = "".join(run_text)
