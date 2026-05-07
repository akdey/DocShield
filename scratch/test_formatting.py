import docx
from pathlib import Path
import subprocess

def test_formatting_preservation():
    # 1. Create a formatted test document
    doc = docx.Document()
    p = doc.add_paragraph('Please contact ')
    p.add_run('John Doe').bold = True
    p.add_run(' at his email ')
    p.add_run('john.doe@example.com').italic = True
    p.add_run('.')
    
    test_file = Path('scratch/test_format.docx')
    test_file.parent.mkdir(exist_ok=True)
    doc.save(test_file)
    
    print("Created test document with Bold and Italic runs.")
    
    # 2. Anonymize
    print("\nAnonymizing...")
    subprocess.run(["uv", "run", "docshield", "anonymize", str(test_file)])
    
    # Check anonymized file
    session_dir = list(Path('scratch').glob('*_output'))[0]
    session_name = session_dir.name.replace('_output', '')
    masked_file = session_dir / f"{session_name}_1.docx"
    
    masked_doc = docx.Document(masked_file)
    print("\nMasked Document Runs:")
    for para in masked_doc.paragraphs:
        for run in para.runs:
            print(f"- Run: '{run.text}' (Bold: {run.bold}, Italic: {run.italic})")
            
    # 3. Deanonymize
    print("\nDeanonymizing...")
    subprocess.run(["uv", "run", "docshield", "deanonymize", str(session_dir)])
    
    restored_file = session_dir / "restored" / f"restored_{session_name}_1.docx"
    restored_doc = docx.Document(restored_file)
    
    print("\nRestored Document Runs:")
    for para in restored_doc.paragraphs:
        for run in para.runs:
            print(f"- Run: '{run.text}' (Bold: {run.bold}, Italic: {run.italic})")
            if "John Doe" in run.text:
                assert run.bold == True, "Lost bold formatting!"
            if "john.doe" in run.text:
                assert run.italic == True, "Lost italic formatting!"

    print("\nSUCCESS! Formatting was preserved.")
    
if __name__ == "__main__":
    # Clean up old test data
    import shutil
    for d in Path('scratch').glob('*_output'):
        shutil.rmtree(d)
        
    test_formatting_preservation()
